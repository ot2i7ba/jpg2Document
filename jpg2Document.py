#!/usr/bin/env python3
# jpg2Document.py
# Copyright (c) 2024 ot2i7ba
# https://github.com/ot2i7ba/
# This code is licensed under the MIT License (see LICENSE for details).

"""
jpg2Document.py v0.1.0 - 2026-04-20
Generate a Word document from a template by inserting a table of compressed,
landscape-oriented images in place of a specified placeholder.
"""

__version__ = "0.1.0"

import argparse
import itertools
import logging
import os
import sys
import tempfile
import threading
import time
from concurrent.futures import ProcessPoolExecutor, as_completed
from dataclasses import dataclass
from math import ceil
from multiprocessing import freeze_support
from pathlib import Path
from typing import List, Optional, Tuple

from PIL import Image
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.table import Table
from docx.text.paragraph import Paragraph

# ---------------------------------------------------------------------------
# Constants / Defaults
# ---------------------------------------------------------------------------

DEFAULT_TEMPLATE        = "jpg2Document.docx"
DEFAULT_OUTPUT          = "pictures.docx"
DEFAULT_PLACEHOLDER     = "<<jpg2Document>>"
DEFAULT_IMAGE_WIDTH_CM  = 9.2
DEFAULT_GAP_WIDTH_CM    = 0.05
DEFAULT_MAX_IMAGE_WIDTH = 1200
DEFAULT_JPEG_QUALITY    = 80
DEFAULT_DOC_AUTHOR      = "jpg2Document"
DEFAULT_DOC_COMMENTS    = "jpg2Document by ot2i7ba"
DEFAULT_EXTENSIONS      = ".jpg,.jpeg,.png"
DEFAULT_RESAMPLE        = "lanczos"

RESAMPLE_FILTERS = {
    "lanczos":  Image.LANCZOS,
    "bicubic":  Image.BICUBIC,
    "bilinear": Image.BILINEAR,
    "nearest":  Image.NEAREST,
}

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

@dataclass
class Config:
    template:       Path
    output:         Path
    placeholder:    str
    input_dir:      Path
    extensions:     Tuple[str, ...]
    image_width_cm: float
    gap_width_cm:   float
    max_width_px:   int
    jpeg_quality:   int
    author:         str
    comments:       str
    resample:       str
    workers:        int
    force:          bool
    clear:          bool
    verbose:        bool

    def validate(self) -> None:
        errors: List[str] = []

        if not self.template.exists():
            errors.append(f"Template not found: '{self.template}'")
        elif self.template.suffix.lower() != ".docx":
            errors.append(f"Template must be a .docx file: '{self.template}'")

        if not self.input_dir.exists() or not self.input_dir.is_dir():
            errors.append(f"Input directory not found: '{self.input_dir}'")

        if self.output.suffix.lower() != ".docx":
            errors.append(f"Output must be a .docx file: '{self.output}'")
        elif self.output.exists() and not self.force:
            errors.append(
                f"Output already exists: '{self.output}' — use --force to overwrite."
            )

        out_parent = self.output.parent
        if out_parent != Path(".") and not out_parent.exists():
            errors.append(f"Output directory does not exist: '{out_parent}'")

        if not 1 <= self.jpeg_quality <= 95:
            errors.append(f"--jpeg_quality must be 1–95, got {self.jpeg_quality}")
        if self.image_width_cm <= 0:
            errors.append(f"--image_width must be > 0, got {self.image_width_cm}")
        if self.gap_width_cm < 0:
            errors.append(f"--gap_width must be >= 0, got {self.gap_width_cm}")
        if self.max_width_px <= 0:
            errors.append(f"--max_px must be > 0, got {self.max_width_px}")
        if not self.extensions:
            errors.append("At least one file extension must be provided")
        for ext in self.extensions:
            if not ext.startswith("."):
                errors.append(f"Extension must start with '.': '{ext}'")
        if self.resample not in RESAMPLE_FILTERS:
            errors.append(
                f"Unknown resample filter '{self.resample}'. "
                f"Choose from: {', '.join(RESAMPLE_FILTERS)}"
            )
        if self.workers < 1:
            errors.append(f"--workers must be >= 1, got {self.workers}")

        if errors:
            raise ValueError("Configuration errors:\n  " + "\n  ".join(errors))


# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

def setup_logging(verbose: bool) -> None:
    logging.basicConfig(
        format="%(asctime)s %(levelname)s: %(message)s",
        level=logging.DEBUG if verbose else logging.INFO,
        datefmt="%H:%M:%S",
    )


# ---------------------------------------------------------------------------
# Spinner
# ---------------------------------------------------------------------------

class Spinner:
    """Thread-based console spinner; call with silent=True to suppress."""

    _FRAMES = ["-", "/", "|", "\\"]

    def __init__(self, label: str, silent: bool = False) -> None:
        self.label = label
        self.silent = silent
        self._stop = threading.Event()
        self._thread = threading.Thread(target=self._spin, daemon=True)

    def _spin(self) -> None:
        blank = " " * (len(self.label) + 20)
        for frame in itertools.cycle(self._FRAMES):
            if self._stop.is_set():
                break
            sys.stdout.write(f"\rProcessing {self.label} ... {frame}")
            sys.stdout.flush()
            time.sleep(0.1)
        sys.stdout.write(f"\r{blank}\r")
        sys.stdout.flush()

    def start(self) -> "Spinner":
        if not self.silent:
            self._thread.start()
        return self

    def stop(self) -> None:
        if self._thread.is_alive():
            self._stop.set()
            self._thread.join()

    def __enter__(self) -> "Spinner":
        return self.start()

    def __exit__(self, *_) -> None:
        self.stop()


# ---------------------------------------------------------------------------
# Image processing
# ---------------------------------------------------------------------------

def get_image_files(input_dir: Path, extensions: Tuple[str, ...]) -> List[Path]:
    return sorted(
        (p for p in input_dir.iterdir()
         if p.is_file() and p.suffix.lower() in extensions),
        key=lambda p: p.name.lower(),
    )


def _compress_one(
    args: Tuple[Path, Path, int, int, str]
) -> Tuple[Optional[Path], str]:
    """Compress and scale a single image. Module-level for multiprocessing pickling."""
    src, dst, max_px, quality, resample_key = args
    filter_ = RESAMPLE_FILTERS.get(resample_key, Image.LANCZOS)
    try:
        with Image.open(src) as img:
            img.load()
            if img.width <= img.height:
                return None, f"skipped (portrait/square): {src.name}"
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            if img.width > max_px:
                scale = max_px / img.width
                img = img.resize(
                    (max_px, max(1, int(img.height * scale))), filter_
                )
            img.save(dst, "JPEG", optimize=True, quality=quality)
        return dst, f"compressed: {src.name}"
    except Exception as exc:
        return None, f"error ({src.name}): {exc}"


def process_images(
    input_dir: Path,
    out_dir: Path,
    extensions: Tuple[str, ...],
    max_width_px: int,
    jpeg_quality: int,
    resample: str,
    workers: int,
) -> Tuple[List[Path], int, int]:
    """
    Compress all landscape images from input_dir into out_dir.
    Returns (sorted_compressed_paths, n_inserted, n_skipped).
    """
    sources = get_image_files(input_dir, extensions)
    if not sources:
        return [], 0, 0

    tasks = [
        (src, out_dir / f"{i:05d}_{src.stem}.jpg", max_width_px, jpeg_quality, resample)
        for i, src in enumerate(sources)
    ]

    success: set = set()
    skipped = 0

    if workers > 1 and len(tasks) > 1:
        with ProcessPoolExecutor(max_workers=workers) as pool:
            for fut in as_completed(pool.submit(_compress_one, t) for t in tasks):
                path, status = fut.result()
                logging.debug(status)
                if path is not None:
                    success.add(path)
                else:
                    skipped += 1
    else:
        for task in tasks:
            path, status = _compress_one(task)
            logging.debug(status)
            if path is not None:
                success.add(path)
            else:
                skipped += 1

    logging.debug("Compression done: %d ok, %d skipped", len(success), skipped)
    # Preserve original alphabetical sort order
    ordered = [t[1] for t in tasks if t[1] in success]
    return ordered, len(ordered), skipped


# ---------------------------------------------------------------------------
# Document building
# ---------------------------------------------------------------------------

def _remove_table_borders(table: Table) -> None:
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.append(tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def _insert_table_after(
    document: Document, paragraph: Paragraph, rows: int, cols: int
) -> Table:
    tmp = document.add_table(rows=rows, cols=cols)
    tbl_el = tmp._tbl
    tmp._element.getparent().remove(tmp._element)
    paragraph._p.addnext(tbl_el)
    return Table(tbl_el, document)


def build_image_table(
    document: Document,
    paragraph: Paragraph,
    image_paths: List[Path],
    image_width_cm: float,
    gap_width_cm: float,
) -> None:
    total = len(image_paths)
    pairs = ceil(total / 2)
    table = _insert_table_after(document, paragraph, pairs * 2, 3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _remove_table_borders(table)

    col_widths = [image_width_cm, gap_width_cm, image_width_cm]
    for row in table.rows:
        for cell, w in zip(row.cells, col_widths):
            cell.width = Cm(w)

    idx = 0
    for i in range(pairs):
        cells = table.rows[i * 2].cells
        for col in (0, 2):
            if idx < total:
                cells[col].paragraphs[0].add_run().add_picture(
                    str(image_paths[idx]), width=Cm(image_width_cm)
                )
                idx += 1


# ---------------------------------------------------------------------------
# Application
# ---------------------------------------------------------------------------

class JPG2DocumentApp:
    def __init__(self, cfg: Config) -> None:
        self.cfg = cfg

    def run(self) -> None:
        cfg = self.cfg

        candidates = get_image_files(cfg.input_dir, cfg.extensions)
        if not candidates:
            raise RuntimeError(
                f"No image files found in '{cfg.input_dir}' "
                f"(extensions: {', '.join(cfg.extensions)})"
            )
        logging.info(
            "%d candidate image(s) found in '%s'", len(candidates), cfg.input_dir
        )

        logging.info("Loading template '%s'", cfg.template)
        doc = self._load_template()
        placeholder_para = self._find_placeholder(doc)
        logging.info("Placeholder '%s' found", cfg.placeholder)

        with tempfile.TemporaryDirectory() as tmpdir:
            with Spinner("images", silent=cfg.verbose):
                images, inserted, skipped = process_images(
                    cfg.input_dir,
                    Path(tmpdir),
                    cfg.extensions,
                    cfg.max_width_px,
                    cfg.jpeg_quality,
                    cfg.resample,
                    cfg.workers,
                )

            if not images:
                raise RuntimeError(
                    f"No landscape images to insert — {skipped} file(s) skipped "
                    "(portrait, square, or unreadable)"
                )

            logging.info("Inserting %d image(s), %d skipped", inserted, skipped)
            self._insert_images(doc, placeholder_para, images)

            if cfg.clear:
                os.system("cls" if os.name == "nt" else "clear")

            logging.info("Saving '%s'", cfg.output)
            self._save_document(doc)

        print()
        logging.info(
            "Done. %d image(s) inserted, %d skipped → '%s'",
            inserted, skipped, cfg.output,
        )

    def _load_template(self) -> Document:
        try:
            return Document(str(self.cfg.template))
        except Exception as exc:
            raise RuntimeError(
                f"Cannot open template '{self.cfg.template}': {exc}"
            ) from exc

    def _find_placeholder(self, doc: Document) -> Paragraph:
        all_paras: List[Paragraph] = list(doc.paragraphs)
        for sec in doc.sections:
            all_paras.extend(sec.header.paragraphs)
            all_paras.extend(sec.footer.paragraphs)
        for para in all_paras:
            if self.cfg.placeholder in para.text:
                return para
        raise ValueError(
            f"Placeholder '{self.cfg.placeholder}' not found in template. "
            "Ensure it is plain, unsplit text in the document body."
        )

    def _insert_images(
        self, doc: Document, para: Paragraph, images: List[Path]
    ) -> None:
        # Prefer run-level replacement to preserve character formatting.
        # Fall back to para.text if the placeholder spans multiple runs.
        for run in para.runs:
            if self.cfg.placeholder in run.text:
                run.text = run.text.replace(self.cfg.placeholder, "")
                break
        else:
            para.text = para.text.replace(self.cfg.placeholder, "")
        build_image_table(
            doc, para, images, self.cfg.image_width_cm, self.cfg.gap_width_cm
        )

    def _save_document(self, doc: Document) -> None:
        doc.core_properties.author = self.cfg.author
        doc.core_properties.comments = self.cfg.comments
        try:
            doc.save(str(self.cfg.output))
        except OSError as exc:
            raise RuntimeError(f"Cannot save '{self.cfg.output}': {exc}") from exc


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args() -> Config:
    cpu_count = os.cpu_count() or 1
    parser = argparse.ArgumentParser(
        description="Insert compressed landscape images into a Word template.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument("--version", action="version", version=f"%(prog)s {__version__}")

    g_io = parser.add_argument_group("input / output")
    g_io.add_argument("--template",   type=Path, default=Path(DEFAULT_TEMPLATE))
    g_io.add_argument("--output",     type=Path, default=Path(DEFAULT_OUTPUT))
    g_io.add_argument("--input_dir",  type=Path, default=Path.cwd())
    g_io.add_argument("--extensions", type=str,  default=DEFAULT_EXTENSIONS,
                      help="Comma-separated image extensions, e.g. .jpg,.png")
    g_io.add_argument("--placeholder",type=str,  default=DEFAULT_PLACEHOLDER)
    g_io.add_argument("--force", action="store_true",
                      help="Overwrite output file if it already exists")

    g_img = parser.add_argument_group("image options")
    g_img.add_argument("--image_width",  type=float, default=DEFAULT_IMAGE_WIDTH_CM,
                       help="Width (cm) per image in the table")
    g_img.add_argument("--gap_width",    type=float, default=DEFAULT_GAP_WIDTH_CM,
                       help="Gap (cm) between image columns")
    g_img.add_argument("--max_px",       type=int,   default=DEFAULT_MAX_IMAGE_WIDTH,
                       help="Maximum image width in pixels before downscaling")
    g_img.add_argument("--jpeg_quality", type=int,   default=DEFAULT_JPEG_QUALITY,
                       help="JPEG compression quality (1–95)")
    g_img.add_argument("--resample", choices=list(RESAMPLE_FILTERS),
                       default=DEFAULT_RESAMPLE,
                       help="Resampling filter for downscaling")

    g_doc = parser.add_argument_group("document metadata")
    g_doc.add_argument("--doc_author",   type=str, default=DEFAULT_DOC_AUTHOR)
    g_doc.add_argument("--doc_comments", type=str, default=DEFAULT_DOC_COMMENTS)

    g_run = parser.add_argument_group("runtime")
    g_run.add_argument("--workers", type=int, default=cpu_count,
                       help="Parallel worker processes for image compression")
    g_run.add_argument("--clear",   action="store_true",
                       help="Clear console before saving output")
    g_run.add_argument("--verbose", action="store_true",
                       help="Enable debug logging (also suppresses spinner)")

    args = parser.parse_args()
    exts = tuple(
        e.strip().lower() for e in args.extensions.split(",") if e.strip()
    )

    return Config(
        template       = args.template.resolve(),
        output         = args.output,
        placeholder    = args.placeholder,
        input_dir      = args.input_dir.resolve(),
        extensions     = exts,
        image_width_cm = args.image_width,
        gap_width_cm   = args.gap_width,
        max_width_px   = args.max_px,
        jpeg_quality   = args.jpeg_quality,
        author         = args.doc_author,
        comments       = args.doc_comments,
        resample       = args.resample,
        workers        = args.workers,
        force          = args.force,
        clear          = args.clear,
        verbose        = args.verbose,
    )


def main() -> None:
    cfg = parse_args()
    setup_logging(cfg.verbose)
    try:
        cfg.validate()
    except ValueError as exc:
        logging.error("%s", exc)
        sys.exit(2)
    try:
        JPG2DocumentApp(cfg).run()
    except (RuntimeError, ValueError) as exc:
        logging.error("%s", exc)
        sys.exit(1)
    except Exception as exc:
        logging.exception("Unexpected error: %s", exc)
        sys.exit(1)


if __name__ == "__main__":
    freeze_support()
    main()
