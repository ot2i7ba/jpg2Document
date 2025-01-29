# jpg2Dokument.py
# Copyright (c) 2024 ot2i7ba
# https://github.com/ot2i7ba/
# This code is licensed under the MIT License (see LICENSE for details).

"""
jpg2Dokument.py v0.0.1 - 2025-01-29
Generate a Word document from a template, inserting a table of compressed,
landscape-oriented images in place of a given placeholder. If the placeholder
is not found, the script aborts immediately (no images are processed).
"""

# Python standard library imports
import os
import sys
import itertools
import time
import threading
from math import ceil
from typing import List
import shutil
import argparse

# Third-party library imports
from PIL import Image, UnidentifiedImageError
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.table import Table

# Default values
DEFAULT_TEMPLATE = "jpg2Dokument.docx"            # Default Word template name
DEFAULT_OUTPUT = "pictures.docx"                  # Default output (generated) document name
DEFAULT_PLACEHOLDER = "<<jpg2Dokument>>"          # Placeholder text to be replaced in the template
DEFAULT_COMPRESSED_FOLDER = "compressed"          # Folder name to store compressed images
DEFAULT_IMAGE_WIDTH_CM = 9.2                      # Width in centimeters for each inserted image
DEFAULT_GAP_WIDTH_CM = 0.05                       # Gap in centimeters between image columns
DEFAULT_MAX_IMAGE_WIDTH_PX = 1200                 # Maximum width in pixels before scaling images down
DEFAULT_JPEG_QUALITY = 80                         # JPEG compression quality (0-100)
DEFAULT_DOC_AUTHOR = "jpg2Dokument"               # Document property: author
DEFAULT_DOC_COMMENTS = "jpg2Dokument by ot2i7ba"  # Document property: comments

def clear_console() -> None:
    """Clear the console screen."""
    if os.name == 'nt':
        os.system('cls')
    else:
        os.system('clear')

def print_blank_line() -> None:
    """Print a blank line."""
    print()

def spinner(stop_event: threading.Event, task_name: str, lock: threading.Lock) -> None:
    """
    Displays a simple rotating spinner while 'stop_event' is not set.
    Used to indicate progress for a specific 'task_name'.
    """
    spinner_cycle = itertools.cycle(['-', '/', '-', '\\'])
    while not stop_event.is_set():
        with lock:
            sys.stdout.write(f'\rProcessing {task_name} ... {next(spinner_cycle)}')
            sys.stdout.flush()
        time.sleep(0.1)
    # After stop_event is set, clear the spinner line
    with lock:
        sys.stdout.write('\r')
        sys.stdout.flush()

def remove_table_borders(table: Table) -> None:
    """Remove all visible borders from the given docx 'table'."""
    tbl = table._tbl
    if tbl.tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    else:
        tblPr = tbl.tblPr

    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'none')
        tblBorders.append(edge_el)
    tblPr.append(tblBorders)

def compress_and_scale_image(
    input_path: str,
    output_path: str,
    max_width_px: int,
    quality: int
) -> bool:
    """
    Attempt to open 'input_path' as an image, resize it to 'max_width_px'
    if wider, and save it as a JPEG to 'output_path' with the specified 'quality'.

    Returns:
        True if the image was successfully opened and saved, False otherwise.
    """
    try:
        with Image.open(input_path) as img:
            # Only accept landscape images (width > height).
            if img.width <= img.height:
                return False

            if img.width > max_width_px:
                ratio = max_width_px / float(img.width)
                new_height = int(img.height * ratio)
                img = img.resize((max_width_px, new_height), Image.LANCZOS)

            img.save(output_path, 'JPEG', optimize=True, quality=quality)

        return True

    except (OSError, UnidentifiedImageError) as err:
        print(f"\nSkipping '{input_path}': could not open or process the image ({err}).")
        return False

def insert_table_after_paragraph(
    document: Document,
    paragraph,
    rows: int,
    cols: int
) -> Table:
    """
    Create a new table with (rows x cols) cells, then insert it immediately
    after the given paragraph in the docx 'document'.

    Returns:
        A docx.table.Table object representing the newly inserted table.
    """
    tmp_table = document.add_table(rows=rows, cols=cols)
    tbl_element = tmp_table._tbl
    tmp_table._element.getparent().remove(tmp_table._element)
    paragraph._p.addnext(tbl_element)

    return Table(tbl_element, paragraph._parent)

def create_table_with_images(
    document: Document,
    paragraph,
    image_paths: List[str],
    image_width_cm: float,
    gap_width_cm: float
) -> None:
    """
    Build a 3-column table of images (left image, narrow gap, right image),
    inserting the table immediately after 'paragraph'. Each pair of images is
    followed by an empty row. 'image_paths' should contain paths to valid images.

    Args:
        document: The Word document object.
        paragraph: The paragraph after which the table should be inserted.
        image_paths: List of image file paths to insert.
        image_width_cm: Width (in cm) for each inserted image.
        gap_width_cm: Width (in cm) for the gap column.
    """
    total_images = len(image_paths)
    pairs = ceil(total_images / 2)
    rows_needed = pairs * 2
    cols = 3

    table = insert_table_after_paragraph(document, paragraph, rows_needed, cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    remove_table_borders(table)

    for row in table.rows:
        row.cells[0].width = Cm(image_width_cm)
        row.cells[1].width = Cm(gap_width_cm)
        row.cells[2].width = Cm(image_width_cm)

    idx_image = 0
    for pair_idx in range(pairs):
        row_images = table.rows[pair_idx * 2]
        row_empty = table.rows[pair_idx * 2 + 1]  # stays empty

        if idx_image < total_images:
            cell_left = row_images.cells[0]
            p_left = cell_left.paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_left.add_run().add_picture(image_paths[idx_image], width=Cm(image_width_cm))
            idx_image += 1

        if idx_image < total_images:
            cell_right = row_images.cells[2]
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.add_run().add_picture(image_paths[idx_image], width=Cm(image_width_cm))
            idx_image += 1

def parse_arguments() -> argparse.Namespace:
    """
    Parse optional command-line arguments. If none are provided,
    default values are used.
    """
    parser = argparse.ArgumentParser(
        description="Insert compressed landscape images into a Word template at a specified placeholder."
    )

    parser.add_argument(
        "--template",
        type=str,
        default=None,
        help="Path to the Word template (.docx)."
    )
    parser.add_argument(
        "--output",
        type=str,
        default=None,
        help="Path for the output Word file."
    )
    parser.add_argument(
        "--placeholder",
        type=str,
        default=None,
        help="Placeholder text in the template to be replaced."
    )
    parser.add_argument(
        "--image_width",
        type=float,
        default=None,
        help="Width in cm for each inserted image."
    )
    parser.add_argument(
        "--gap_width",
        type=float,
        default=None,
        help="Width in cm for the gap between images."
    )
    parser.add_argument(
        "--max_px",
        type=int,
        default=None,
        help="Maximum width in pixels before image is scaled down."
    )
    parser.add_argument(
        "--jpeg_quality",
        type=int,
        default=None,
        help="JPEG compression quality (0-100)."
    )
    parser.add_argument(
        "--doc_author",
        type=str,
        default=None,
        help="Value for the document's author property."
    )
    parser.add_argument(
        "--doc_comments",
        type=str,
        default=None,
        help="Value for the document's comments property."
    )

    return parser.parse_args()

def main() -> None:
    """
    Main flow:
      1) Read defaults and possibly override them from CLI arguments.
      2) Open the template.
      3) Check if the placeholder is found; if not, abort immediately.
      4) Collect and compress all landscape images in the current directory.
         (Here we use a spinner to show progress.)
      5) Insert them into the template in place of the placeholder.
      6) Set document properties and save.
      7) Clean up temporary files.
    """
    args = parse_arguments()

    docx_template = args.template or DEFAULT_TEMPLATE
    docx_output = args.output or DEFAULT_OUTPUT
    placeholder_text = args.placeholder or DEFAULT_PLACEHOLDER

    compressed_folder = DEFAULT_COMPRESSED_FOLDER
    image_width_cm = args.image_width or DEFAULT_IMAGE_WIDTH_CM
    gap_width_cm = args.gap_width or DEFAULT_GAP_WIDTH_CM
    max_image_width_px = args.max_px or DEFAULT_MAX_IMAGE_WIDTH_PX
    jpeg_quality = args.jpeg_quality or DEFAULT_JPEG_QUALITY

    doc_author = args.doc_author or DEFAULT_DOC_AUTHOR
    doc_comments = args.doc_comments or DEFAULT_DOC_COMMENTS

    if not os.path.exists(docx_template):
        print(f"Template '{docx_template}' was not found. Aborting.")
        return

    # Attempt to open the template
    try:
        document = Document(docx_template)
    except Exception as err:
        print(f"Failed to open template '{docx_template}': {err}")
        return

    # -------- NEW STEP: Check if placeholder is present BEFORE processing images.
    placeholder_found = any(
        placeholder_text in paragraph.text for paragraph in document.paragraphs
    )
    if not placeholder_found:
        print(f"Placeholder '{placeholder_text}' was not found. Aborting.")
        return

    # If we reach here, we do have the placeholder -> proceed.
    current_dir = os.getcwd()
    valid_ext = (".jpg", ".jpeg", ".png")

    image_files = [
        f for f in os.listdir(current_dir) if f.lower().endswith(valid_ext)
    ]
    image_files.sort(key=lambda x: x.lower())

    if not image_files:
        print("No image files found. Aborting.")
        return

    compressed_dir = os.path.join(current_dir, compressed_folder)
    os.makedirs(compressed_dir, exist_ok=True)

    stop_event = threading.Event()
    lock = threading.Lock()
    spinner_thread = threading.Thread(
        target=spinner,
        args=(stop_event, "images", lock),
        daemon=True
    )
    spinner_thread.start()

    compressed_paths: List[str] = []
    for i, filename in enumerate(image_files):
        in_path = os.path.join(current_dir, filename)
        out_path = os.path.join(compressed_dir, f"compressed_{i}.jpg")
        success = compress_and_scale_image(in_path, out_path, max_image_width_px, jpeg_quality)
        if success:
            compressed_paths.append(out_path)

    # End the spinner
    stop_event.set()
    spinner_thread.join()
    print()

    if not compressed_paths:
        print("No valid landscape images could be processed. Aborting.")
        shutil.rmtree(compressed_dir, ignore_errors=True)
        return

    # Now replace the placeholder and create the table
    for paragraph in document.paragraphs:
        if placeholder_text in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder_text, "")
            create_table_with_images(document, paragraph, compressed_paths, image_width_cm, gap_width_cm)
            break

    print_blank_line()
    print(f"Placeholder '{placeholder_text}' found and replaced.")

    # Set document properties
    document.core_properties.author = doc_author
    document.core_properties.comments = doc_comments

    document.save(docx_output)
    print(f"Document saved as: {docx_output}")

    shutil.rmtree(compressed_dir, ignore_errors=True)
    print(f"Temporary folder '{compressed_folder}' removed.")

    print_blank_line()
    print("Time saved once again. Thanks, ot2i7ba!")
    print("Check the result and adjust it if necessary!")

if __name__ == "__main__":
    clear_console()
    try:
        main()
    except KeyboardInterrupt:
        print_blank_line()
        print("\nProcess interrupted by user. Exiting gracefully...")
        sys.exit(0)
